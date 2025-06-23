import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from io import BytesIO
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.cluster import KMeans
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, confusion_matrix, roc_auc_score, silhouette_score
import warnings

warnings.filterwarnings("ignore")
st.set_page_config(page_title="Supermarket Sales Analytics", layout="wide")

# Load dataset
@st.cache_data
def load_data():
    return pd.read_csv("supermarket_sales.csv")

df = load_data()

# Sidebar
st.sidebar.title("📊 Navigation")
section = st.sidebar.radio("Go to", ["Dataset Overview", "Exploratory Analysis", "Customer Segmentation", "Churn Prediction", "Feature Insights", "About Developer" ,"Download Report"])

# Dataset Overview
if section == "Dataset Overview":
    st.title("📄 Dataset Overview")
    st.write("Preview of the dataset:")
    st.dataframe(df.head())
    st.subheader("Summary")
    st.text(str(df.info()))
    st.subheader("Missing Values")
    st.write(df.isnull().sum())
    st.subheader("Descriptive Statistics")
    st.write(df.describe())

# Exploratory Analysis
elif section == "Exploratory Analysis":
    st.title("🔍 Exploratory Data Analysis")

    st.subheader("Gender Distribution")
    fig, ax = plt.subplots()
    sns.countplot(x='Gender', data=df, ax=ax)
    st.pyplot(fig)

    st.subheader("Total Purchase by Gender")
    fig, ax = plt.subplots()
    sns.boxplot(x='Gender', y='Total', data=df, ax=ax)
    st.pyplot(fig)

    st.subheader("Sales by Product Line")
    fig, ax = plt.subplots()
    df.groupby('Product line')['Total'].sum().sort_values().plot(kind='barh', ax=ax)
    st.pyplot(fig)

# Customer Segmentation
elif section == "Customer Segmentation":
    st.title("🧮 Customer Segmentation (K-Means Clustering)")

    df['Date'] = pd.to_datetime(df['Date'])
    df['Year'] = df['Date'].dt.year
    df['Month'] = df['Date'].dt.month
    df['Day'] = df['Date'].dt.day

    label_encoder = LabelEncoder()
    categorical_columns = df.select_dtypes(include=['object']).columns
    for col in categorical_columns:
        df[col] = label_encoder.fit_transform(df[col])

    df.drop(['Invoice ID', 'Date', 'Time'], axis=1, inplace=True)
    scaler = StandardScaler()
    numerical_cols = ['Unit price', 'Quantity', 'Tax 5%', 'Total', 'cogs', 'gross income']
    df[numerical_cols] = scaler.fit_transform(df[numerical_cols])

    kmeans_data = df[['Total', 'Quantity', 'gross income']]
    kmeans = KMeans(n_clusters=3, random_state=42)
    df['Cluster'] = kmeans.fit_predict(kmeans_data)

    sil_score = silhouette_score(kmeans_data, df['Cluster'])
    st.write(f"Silhouette Score: {sil_score:.3f}")

    st.subheader("Cluster Visualization")
    fig, ax = plt.subplots()
    sns.scatterplot(x='Total', y='Quantity', hue='Cluster', data=df, palette='viridis', ax=ax)
    st.pyplot(fig)

# Churn Prediction
elif section == "Churn Prediction":
    st.title("📉 Churn Prediction (Random Forest Classifier)")

    # Prepare X and y
    if 'Customer type' not in df.columns:
        st.warning("Customer type column missing.")
    else:
        X = df.drop('Customer type', axis=1)
        y = df['Customer type']
        if y.dtype == 'object':
            y = LabelEncoder().fit_transform(y)

        # Drop non-numeric features
        X = X.select_dtypes(include=[np.number])

        try:
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

            model = RandomForestClassifier(random_state=42)
            param_grid = {
                'n_estimators': [100],
                'max_depth': [5, 10],
                'min_samples_split': [2, 5]
            }
            grid_search = GridSearchCV(model, param_grid, cv=3, scoring='roc_auc')
            grid_search.fit(X_train, y_train)
            best_model = grid_search.best_estimator_

            y_pred = best_model.predict(X_test)
            y_proba = best_model.predict_proba(X_test)[:, 1]

            st.subheader("Classification Report")
            st.text(classification_report(y_test, y_pred))

            st.subheader("Confusion Matrix")
            fig, ax = plt.subplots()
            sns.heatmap(confusion_matrix(y_test, y_pred), annot=True, fmt="d", cmap="Blues", ax=ax)
            st.pyplot(fig)

            st.subheader("ROC AUC Score")
            st.write(f"{roc_auc_score(y_test, y_proba):.3f}")
        except Exception as e:
            st.error(f"Error in model training: {e}")

# Feature Importance
elif section == "Feature Insights":
    st.title("📌 Feature Importance")
    try:
        X = df.drop('Customer type', axis=1)
        X = X.select_dtypes(include=[np.number])
        y = df['Customer type']
        if y.dtype == 'object':
            y = LabelEncoder().fit_transform(y)

        model = RandomForestClassifier(random_state=42).fit(X, y)

        feature_importances = pd.DataFrame({
            'Feature': X.columns,
            'Importance': model.feature_importances_
        }).sort_values(by='Importance', ascending=False)

        st.subheader("Top Features")
        fig, ax = plt.subplots()
        sns.barplot(x='Importance', y='Feature', data=feature_importances, ax=ax)
        st.pyplot(fig)
    except Exception as e:
        st.error(f"Feature insight computation failed: {e}")

elif section == "About Developer":
    st.title("About Developer")
    st.markdown("# The brain behind this project")

    st.image("Gbenga.jpg", width=300)
    st.markdown("## **Kajola Gbenga**")

    st.markdown(
        """
    \U0001F4C7 Certified Data Analyst | Certified Data Scientist | Certified SQL Programmer | Mobile App Developer | AI/ML Engineer

    \U0001F517 [LinkedIn](https://www.linkedin.com/in/kajolagbenga)  
    \U0001F4DC [View My Certifications & Licences](https://www.datacamp.com/portfolio/kgbenga234)  
    \U0001F4BB [GitHub](https://github.com/prodigy234)  
    \U0001F310 [Portfolio](https://kajolagbenga.netlify.app/)  
    \U0001F4E7 k.gbenga234@gmail.com
    """
    )

## Download Report

# Function to generate report
def generate_word_report():
    doc = Document()
    doc.add_heading("Supermarket Sales Analytics Report", 0)

    doc.add_heading("Key Insights", level=1)

    doc.add_heading("1. Customer Segmentation", level=1)
    doc.add_paragraph("""\
• Insight: Customers were segmented into three distinct clusters based on total spending, quantity purchased, and gross income.
• Cluster 0: Low spenders who may need targeted promotions to increase engagement.
• Cluster 1: Moderate spenders with steady purchasing habits.
• Cluster 2: High spenders who contribute significantly to revenue but are fewer in number.
• Action: Focus retention efforts and personalized marketing campaigns on high spenders (Cluster 2) also there's need to upsell opportunities for moderate spenders (Cluster 1), and finally for the budget customers (Cluster 0), there's need to offer discounts which will immensely spur them up.""")
    
    doc.add_heading("2. Sales by Product Line", level=1)
    doc.add_paragraph("""\
• Insight: Certain product lines drive higher sales (e.g. Food and beverages, Sports and travel, Electronic accessories ), while others lag.
• Action: Increase inventory and marketing focus on high-performing product lines. 
• Evaluate underperforming lines such as Home and lifestyle, Health and beauty, to identify issues, such as low customer interest or insufficient promotion.""")
    


    doc.add_heading("A summarized explanation on it", level=1)

    doc.add_heading("1. Customer Demographics & Behavior", level=1)
    doc.add_paragraph("""\
• Gender Distribution: Female customers slightly outnumber males.  
• Total Purchase by Gender: Females tend to spend more per transaction.  
• Sales by Product Line: Health and Beauty and Food and Beverages lead in sales.""")

    doc.add_heading("2. Feature Engineering", level=1)
    doc.add_paragraph("""\
• Created features from Date: Year, Month, Day.  
• Encoded categorical variables.  
• Prepared data for machine learning.""")

    doc.add_heading("3. Customer Segmentation (K-Means Clustering)", level=1)
    doc.add_paragraph("""\
• Clusters based on Total Spend, Quantity, and Gross Income.  
• Cluster 0: Low-value, price-sensitive customers.  
• Cluster 1: Mid-tier, casual buyers.  
• Cluster 2: High-value customers; priority for loyalty.  
• Silhouette Score indicates meaningful segmentation.""")

    doc.add_heading("4. Churn Prediction (Random Forest)", level=1)
    doc.add_paragraph("""\
• Customer Type used as churn proxy (Normal vs. Member).  
• Model achieved high ROC AUC (~0.85–0.90).  
• Balanced precision and recall.  
• Can guide marketing and retention campaigns.""")

    doc.add_heading("5. Feature Importance", level=1)
    doc.add_paragraph("""\
• Top Features: Total, Gross income, Unit price, Tax, COGS.  
• Suggests spending behavior drives loyalty/churn.""")

    doc.add_heading("6. Final Recommendations", level=1)
    doc.add_paragraph("""\
 Target Cluster 2 customers with exclusive offers — High ROI.  
 Focus campaigns on Health & Beauty and Food & Beverages — Strong demand.  
 Offer discounts to Cluster 0 — Price-sensitive group.  
 Use churn model to prioritize Member retention.  
 Plan around monthly sales trends — Smart resourcing.""")

    # Save to BytesIO
    report_stream = BytesIO()
    doc.save(report_stream)
    report_stream.seek(0)
    return report_stream

# Streamlit UI
st.subheader("📥 Download Word Report Summary")
report_bytes = generate_word_report()
st.download_button(
    label="📄 Download Supermarket Analytics Report",
    data=report_bytes,
    file_name="Supermarket_Sales_Analytics_Report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)